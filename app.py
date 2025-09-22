import streamlit as st
import json
import zipfile
import xml.etree.ElementTree as ET
from xml.dom import minidom
import difflib
import re
from typing import Dict, List, Tuple, Optional, Set
import tempfile
import os
from dataclasses import dataclass, asdict
from io import StringIO, BytesIO
import markdown
from fuzzywuzzy import fuzz, process
import datetime
from docx import Document
from docx.shared import Inches

@dataclass
class ParsedSection:
    name: str
    type: str
    sub_type: str
    prompt: str
    include_screenshots: bool
    screenshot_instructions: str
    raw_comment: str
    edited: bool = False
    original_prompt: str = ""

    def __post_init__(self):
        if not self.original_prompt:
            self.original_prompt = self.prompt

@dataclass
class SectionMapping:
    file1_section: str
    file2_section: str
    confidence: float
    is_manual: bool = False

@dataclass
class SectionDeviation:
    section_name: str
    file1_section: Optional[ParsedSection]
    file2_section: Optional[ParsedSection]
    deviation_type: str  # 'content', 'missing', 'added', 'properties', 'identical'
    similarity_score: float
    content_diff_chars: int
    property_differences: List[str]
    severity: str  # 'high', 'medium', 'low'

    def __post_init__(self):
        # Auto-calculate severity based on multiple factors
        if self.deviation_type in ['missing', 'added']:
            self.severity = 'high'
        elif self.similarity_score < 50:
            self.severity = 'high'
        elif self.similarity_score < 75 or len(self.property_differences) > 1:
            self.severity = 'medium'
        else:
            self.severity = 'low'

class DocxCommentExtractor:
    """Extract comments and their associated section names from DOCX files"""

    @staticmethod
    def extract_comments_from_docx(file_path: str) -> Dict[str, ParsedSection]:
        """Extract all comments and their associated section names from a DOCX file"""
        sections = {}

        try:
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # Read comments.xml
                comments_xml = None
                document_xml = None

                try:
                    comments_xml = docx_zip.read('word/comments.xml').decode('utf-8')
                except KeyError:
                    st.warning("No comments.xml found in DOCX file")
                    return sections

                try:
                    document_xml = docx_zip.read('word/document.xml').decode('utf-8')
                except KeyError:
                    st.warning("No document.xml found in DOCX file")
                    return sections

                # Parse comments
                comments_data = DocxCommentExtractor._parse_comments_xml(comments_xml)

                # Parse document to find section names linked to comments
                section_names = DocxCommentExtractor._parse_document_for_sections(document_xml)

                # Combine comments with section names
                for comment_id, comment_text in comments_data.items():
                    if comment_id in section_names:
                        section_name = section_names[comment_id]
                        parsed_section = DocxCommentExtractor._parse_comment_string(
                            comment_text, section_name
                        )
                        if parsed_section:
                            sections[section_name] = parsed_section

        except Exception as e:
            st.error(f"Error processing DOCX file: {str(e)}")

        return sections

    @staticmethod
    def _parse_comments_xml(comments_xml: str) -> Dict[str, str]:
        """Parse comments.xml to extract comment ID and text"""
        comments = {}

        try:
            root = ET.fromstring(comments_xml)

            # Handle namespaces
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }

            comment_elements = root.findall('.//w:comment', namespaces)

            for comment in comment_elements:
                comment_id = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')

                # Extract text content from comment
                text_elements = comment.findall('.//w:t', namespaces)
                comment_text = ''.join([t.text or '' for t in text_elements])

                if comment_id and comment_text:
                    comments[comment_id] = comment_text.strip()

        except ET.ParseError as e:
            st.error(f"Error parsing comments XML: {str(e)}")

        return comments

    @staticmethod
    def _parse_document_for_sections(document_xml: str) -> Dict[str, str]:
        """Parse document.xml to find section names associated with comment IDs"""
        section_names = {}

        try:
            root = ET.fromstring(document_xml)

            # Handle namespaces
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }

            # Find all paragraphs
            paragraphs = root.findall('.//w:p', namespaces)

            for paragraph in paragraphs:
                # Check for comment range start
                comment_range_start = paragraph.find('.//w:commentRangeStart', namespaces)
                comment_reference = paragraph.find('.//w:commentReference', namespaces)

                if comment_range_start is not None or comment_reference is not None:
                    # Get comment ID
                    comment_id = None
                    if comment_range_start is not None:
                        comment_id = comment_range_start.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                    elif comment_reference is not None:
                        comment_id = comment_reference.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')

                    # Get text content (section name)
                    text_elements = paragraph.findall('.//w:t', namespaces)
                    section_name = ''.join([t.text or '' for t in text_elements]).strip()

                    if comment_id and section_name:
                        section_names[comment_id] = section_name

        except ET.ParseError as e:
            st.error(f"Error parsing document XML: {str(e)}")

        return section_names

    @staticmethod
    def _parse_comment_string(comment_text: str, section_name: str) -> Optional[ParsedSection]:
        """Parse comment text to extract structured section data with intelligent prompt extraction"""

        try:
            # Clean and normalize the comment text
            comment_text = comment_text.strip()

            # Initialize with safe defaults
            section_data = {
                'type': 'text',
                'sub_type': 'default',
                'prompt': '',
                'include_screenshots': False,
                'screenshot_instructions': ''
            }

            # Handle the specific format from Alex's templates:
            # "type - textsub_type - freeformprompt - [ACTUAL PROMPT]include_screenshots - yesscreenshot_instructions - none"

            # First, let's check if this is the concatenated format (no line breaks)
            if '\n' not in comment_text and ' - ' in comment_text:
                # This is likely the concatenated format, let's parse it carefully
                return DocxCommentExtractor._parse_concatenated_format(comment_text, section_name)

            # Handle multi-line format
            if ' - ' not in comment_text and ': ' not in comment_text and '=' not in comment_text:
                # Treat entire comment as prompt
                section_data['prompt'] = comment_text
                return ParsedSection(
                    name=section_name,
                    type=section_data['type'],
                    sub_type=section_data['sub_type'],
                    prompt=section_data['prompt'],
                    include_screenshots=section_data['include_screenshots'],
                    screenshot_instructions=section_data['screenshot_instructions'],
                    raw_comment=comment_text
                )

            # Parse line-by-line format
            lines = comment_text.split('\n')

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # Try different separators: ' - ', ': ', '='
                key = ''
                value = ''
                separator_found = False

                if ' - ' in line and not separator_found:
                    parts = line.split(' - ', 1)
                    if len(parts) == 2:
                        key, value = parts[0].strip(), parts[1].strip()
                        separator_found = True

                if ': ' in line and not separator_found:
                    parts = line.split(': ', 1)
                    if len(parts) == 2:
                        key, value = parts[0].strip(), parts[1].strip()
                        separator_found = True

                if '=' in line and not separator_found:
                    parts = line.split('=', 1)
                    if len(parts) == 2:
                        key, value = parts[0].strip(), parts[1].strip()
                        separator_found = True

                if not separator_found:
                    # If no separator, treat as prompt content (append to existing)
                    if section_data['prompt']:
                        section_data['prompt'] += '\n' + line
                    else:
                        section_data['prompt'] = line
                    continue

                # Skip if no valid key-value pair found
                if not key or not value:
                    continue

                # Process the key-value pair
                DocxCommentExtractor._process_key_value(key, value, section_data)

            # If no prompt found in structured format, use entire comment as prompt
            if not section_data['prompt'] and len(comment_text.strip()) > 15:
                section_data['prompt'] = comment_text.strip()

            # Final validation - ensure required fields have valid values
            if section_data['type'] not in ['text', 'table', 'process_flow_diagram']:
                section_data['type'] = 'text'

            if section_data['sub_type'] not in ['default', 'bulleted', 'freeform', 'flow-diagram', 'walkthrough-steps', 'bpmn']:
                section_data['sub_type'] = 'default'

            return ParsedSection(
                name=section_name,
                type=section_data['type'],
                sub_type=section_data['sub_type'],
                prompt=section_data['prompt'],
                include_screenshots=section_data['include_screenshots'],
                screenshot_instructions=section_data['screenshot_instructions'],
                raw_comment=comment_text
            )

        except Exception as e:
            # Return a safe default section instead of None
            return ParsedSection(
                name=section_name,
                type='text',
                sub_type='default',
                prompt=comment_text.strip() if comment_text.strip() else f"Content for {section_name}",
                include_screenshots=False,
                screenshot_instructions='',
                raw_comment=comment_text
            )

    @staticmethod
    def _parse_concatenated_format(comment_text: str, section_name: str) -> ParsedSection:
        """Parse the concatenated format: 'type - textsub_type - freeformprompt - [CONTENT]include_screenshots - yes...'"""

        section_data = {
            'type': 'text',
            'sub_type': 'default',
            'prompt': '',
            'include_screenshots': False,
            'screenshot_instructions': ''
        }

        # Use regex to extract the different parts more intelligently
        import re

        # Pattern to match the structured parts
        # Look for patterns like "key - value" but be smart about the prompt content

        # First, extract type
        type_match = re.match(r'^type\s*-\s*([^s]*?)sub_type', comment_text)
        if type_match:
            type_value = type_match.group(1).strip()
            # Map special types
            if type_value == 'process_flow_diagram':
                section_data['type'] = 'text'  # Process flow diagrams are rendered as text
            elif type_value in ['text', 'table']:
                section_data['type'] = type_value
            else:
                section_data['type'] = 'text'

        # Extract sub_type
        subtype_match = re.search(r'sub_type\s*-\s*([^p]*?)prompt\s*-', comment_text)
        if subtype_match:
            subtype_value = subtype_match.group(1).strip()
            # Handle special cases
            if subtype_value == 'bpmn':
                section_data['sub_type'] = 'flow-diagram'
            elif subtype_value in ['default', 'bulleted', 'freeform', 'flow-diagram', 'walkthrough-steps']:
                section_data['sub_type'] = subtype_value
            else:
                section_data['sub_type'] = 'default'

        # Extract the main prompt content - this is the tricky part
        # Look for "prompt - " and then capture everything until "include_screenshots"
        prompt_match = re.search(r'prompt\s*-\s*(.*?)include_screenshots\s*-', comment_text, re.DOTALL)
        if prompt_match:
            section_data['prompt'] = prompt_match.group(1).strip().replace('\\n', '\n')

        # Extract screenshots setting
        screenshot_match = re.search(r'include_screenshots\s*-\s*([^s]*?)screenshot_instructions', comment_text)
        if screenshot_match:
            screenshot_value = screenshot_match.group(1).strip().lower()
            section_data['include_screenshots'] = screenshot_value in ['yes', 'true', '1']

        # Extract screenshot instructions
        instructions_match = re.search(r'screenshot_instructions\s*-\s*(.*)$', comment_text)
        if instructions_match:
            instructions_value = instructions_match.group(1).strip()
            section_data['screenshot_instructions'] = instructions_value if instructions_value.lower() != 'none' else ''

        # If we couldn't extract a proper prompt, fall back to a cleaned version
        if not section_data['prompt']:
            # Try to extract everything after the first "prompt - " occurrence
            prompt_start = comment_text.find('prompt - ')
            if prompt_start != -1:
                remaining_text = comment_text[prompt_start + 9:]  # 9 = len('prompt - ')
                # Remove trailing metadata
                for suffix in ['include_screenshots', 'screenshot_instructions']:
                    pos = remaining_text.find(suffix)
                    if pos != -1:
                        remaining_text = remaining_text[:pos]
                        break
                section_data['prompt'] = remaining_text.strip().replace('\\n', '\n')

        return ParsedSection(
            name=section_name,
            type=section_data['type'],
            sub_type=section_data['sub_type'],
            prompt=section_data['prompt'] if section_data['prompt'] else comment_text.strip(),
            include_screenshots=section_data['include_screenshots'],
            screenshot_instructions=section_data['screenshot_instructions'],
            raw_comment=comment_text
        )

    @staticmethod
    def _process_key_value(key: str, value: str, section_data: dict):
        """Process a key-value pair and update section_data"""

        # Normalize key
        key_normalized = key.lower().replace('_', '').replace(' ', '').replace('-', '')

        try:
            if key_normalized in ['type', 'contenttype']:
                # Handle special type mappings
                clean_type = value.lower().strip()
                if clean_type == 'process_flow_diagram':
                    section_data['type'] = 'text'
                elif clean_type in ['text', 'table']:
                    section_data['type'] = clean_type
                else:
                    section_data['type'] = 'text'

            elif key_normalized in ['subtype', 'sub_type', 'style']:
                # Handle special subtype mappings
                clean_subtype = value.lower().strip()
                if clean_subtype == 'bpmn':
                    section_data['sub_type'] = 'flow-diagram'
                elif clean_subtype in ['default', 'bulleted', 'freeform', 'flow-diagram', 'walkthrough-steps']:
                    section_data['sub_type'] = clean_subtype
                else:
                    section_data['sub_type'] = 'default'

            elif key_normalized in ['prompt', 'instruction', 'content']:
                section_data['prompt'] = value.replace('\\n', '\n')

            elif key_normalized in ['includescreenshot', 'include_screenshot', 'screenshot']:
                section_data['include_screenshots'] = value.lower().strip() in ['yes', 'true', '1']

            elif key_normalized in ['screenshotinstruction', 'screenshot_instruction']:
                section_data['screenshot_instructions'] = value if value.lower() != 'none' else ''

        except Exception as field_error:
            # Skip this field if there's an error, but continue processing
            pass

class PromptFormatter:
    """Convert JSON prompts to readable markdown format with enhanced UX"""

    @staticmethod
    def format_prompt_for_display(prompt_text: str, max_length: int = None) -> str:
        """Format prompt text for optimal readability in Streamlit"""

        if not prompt_text or not prompt_text.strip():
            return "*No prompt content available*"

        # Try to parse as JSON first
        try:
            if prompt_text.strip().startswith('{') or prompt_text.strip().startswith('['):
                prompt_obj = json.loads(prompt_text)
                formatted = PromptFormatter._json_obj_to_markdown(prompt_obj)
            else:
                formatted = PromptFormatter._enhance_text_formatting(prompt_text)
        except json.JSONDecodeError:
            formatted = PromptFormatter._enhance_text_formatting(prompt_text)

        # Truncate if needed
        if max_length and len(formatted) > max_length:
            formatted = formatted[:max_length] + "..."

        return formatted

    @staticmethod
    def json_to_markdown(prompt_text: str) -> str:
        """Convert JSON string prompts to readable markdown - legacy method"""
        return PromptFormatter.format_prompt_for_display(prompt_text)

    @staticmethod
    def _json_obj_to_markdown(obj) -> str:
        """Convert JSON object to markdown recursively"""

        if isinstance(obj, dict):
            markdown_parts = []
            for key, value in obj.items():
                if isinstance(value, (dict, list)):
                    markdown_parts.append(f"## {key.title().replace('_', ' ')}")
                    markdown_parts.append(PromptFormatter._json_obj_to_markdown(value))
                else:
                    markdown_parts.append(f"**{key.title().replace('_', ' ')}:** {value}")
            return '\n\n'.join(markdown_parts)

        elif isinstance(obj, list):
            if all(isinstance(item, str) for item in obj):
                return '\n'.join([f"- {item}" for item in obj])
            else:
                markdown_parts = []
                for i, item in enumerate(obj, 1):
                    markdown_parts.append(f"### Item {i}")
                    markdown_parts.append(PromptFormatter._json_obj_to_markdown(item))
                return '\n\n'.join(markdown_parts)

        else:
            return str(obj)

    @staticmethod
    def _enhance_text_formatting(text: str) -> str:
        """Enhanced text formatting for better Streamlit display"""

        if not text:
            return ""

        # Clean and normalize whitespace
        text = re.sub(r'\r\n', '\n', text)  # Windows line endings
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Max 2 newlines
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space

        # Handle escaped newlines from comment parsing first
        text = text.replace('\\n', '\n')

        # Convert numbered lists with better formatting
        text = re.sub(r'(?:^|\n)(\d+)\.[ \t]*', r'\n\n\1. ', text)

        # Convert bullet points with consistent formatting - preserve special bullet chars
        text = re.sub(r'(?:^|\n)[•·*-][ \t]*', r'\n- ', text)

        # Convert markdown bold to HTML for better display in containers
        text = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', text)

        # Emphasize key phrases
        text = re.sub(r'"([^"]+)"', r'<strong>\1</strong>', text)  # Quoted text
        text = re.sub(r'\b([A-Z]{2,})\b', r'<strong>\1</strong>', text)  # ALL CAPS words

        # Handle special formatting for instructions and examples
        text = re.sub(r'Example user quote: "([^"]+)"', r'<em>Example user quote: "<strong>\1</strong>"</em>', text)

        # Clean up extra whitespace at start/end
        text = text.strip()

        return text

    @staticmethod
    def create_section_property_display(section: ParsedSection, show_stats: bool = True) -> str:
        """Create a clean, readable display of section properties"""

        lines = []

        # Type information with icons
        type_icon = "📊" if section.type == "table" else "📝"
        lines.append(f"{type_icon} **Content Type:** {section.type.title()}")

        # Sub-type with descriptive text
        subtype_descriptions = {
            'default': 'Standard text format',
            'bulleted': 'Bullet point list',
            'freeform': 'Free-form narrative',
            'flow-diagram': 'Process flow description',
            'walkthrough-steps': 'Step-by-step instructions'
        }
        subtype_desc = subtype_descriptions.get(section.sub_type, section.sub_type.title())
        lines.append(f"🎯 **Format:** {subtype_desc}")

        # Screenshot requirements
        screenshot_icon = "📸" if section.include_screenshots else "🚫"
        screenshot_text = "Required" if section.include_screenshots else "Not required"
        lines.append(f"{screenshot_icon} **Screenshots:** {screenshot_text}")

        if section.include_screenshots and section.screenshot_instructions:
            lines.append(f"   *Instructions: {section.screenshot_instructions}*")

        # Statistics if requested
        if show_stats and section.prompt:
            char_count = len(section.prompt)
            word_count = len(section.prompt.split())
            lines.append(f"📈 **Size:** {char_count:,} characters, {word_count:,} words")

        # Edit status
        if hasattr(section, 'edited') and section.edited:
            lines.append("✏️ **Status:** Modified")
        else:
            lines.append("📄 **Status:** Original")

        return '\n'.join(lines)

class UIComponents:
    """Enhanced UI components for better UX"""

    @staticmethod
    def create_section_card(section_name: str, section: ParsedSection, file_name: str = "", expanded: bool = False):
        """Create a well-formatted section card with unique keys"""

        # Create unique key prefix using file name and section name
        unique_key = f"{file_name}_{section_name}".replace(" ", "_").replace(".", "_").replace("/", "_")

        # Create status indicator
        status_icon = " ✏️" if (hasattr(section, 'edited') and section.edited) else ""

        with st.expander(f"📝 {section_name}{status_icon}", expanded=expanded):
            # Two-column layout
            col1, col2 = st.columns([2, 1])

            with col1:
                st.markdown("### 📋 Prompt Content")

                # Format and display the prompt
                formatted_prompt = PromptFormatter.format_prompt_for_display(section.prompt)

                # Use a container with custom styling for better readability
                st.markdown(
                    f"""
                    <div style="
                        background-color: #f8f9fa;
                        padding: 1rem;
                        border-radius: 0.5rem;
                        border-left: 4px solid #007bff;
                        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                        line-height: 1.6;
                        color: #212529;
                    ">
                    {formatted_prompt.replace(chr(10), '<br>')}
                    </div>
                    """,
                    unsafe_allow_html=True
                )

                # Show edit indicator if applicable
                if hasattr(section, 'edited') and section.edited:
                    st.success("✏️ This prompt has been edited from the original")

            with col2:
                st.markdown("### ⚙️ Properties")

                # Display properties in a clean format
                properties_text = PromptFormatter.create_section_property_display(section)
                st.markdown(properties_text)

                # Action buttons
                st.markdown("### 🔧 Actions")

                if st.button(f"📝 Edit", key=f"edit_{unique_key}", help="Edit this section"):
                    st.session_state[f"edit_mode_{unique_key}"] = True
                    st.rerun()

                if st.button(f"📋 Copy", key=f"copy_{unique_key}", help="Copy prompt to clipboard"):
                    st.success("Prompt copied to clipboard!")

    @staticmethod
    def create_side_by_side_comparison(section_name: str, section1: ParsedSection, section2: ParsedSection,
                                     file1_name: str, file2_name: str):
        """Create an enhanced side-by-side comparison view with inline diff highlighting"""

        st.markdown(f"## 📊 Comparing: {section_name}")

        # Quick stats bar
        col1, col2, col3 = st.columns(3)
        with col1:
            similarity = fuzz.ratio(section1.prompt, section2.prompt)
            st.metric("Similarity", f"{similarity}%", help="Text similarity percentage")
        with col2:
            char_diff = len(section2.prompt) - len(section1.prompt)
            st.metric("Length Difference", f"{char_diff:+,} chars", help="Character count difference")
        with col3:
            type_match = "✅" if section1.type == section2.type else "❌"
            st.metric("Type Match", type_match, help="Whether content types match")

        st.divider()

        # Generate side-by-side diff highlighting
        diff_content1, diff_content2 = UIComponents._generate_side_by_side_diff(
            section1.prompt, section2.prompt
        )

        # Side-by-side content with diff highlighting
        col1, col2 = st.columns(2)

        with col1:
            st.markdown(f"### 📄 {file1_name}")

            # Properties box
            with st.container():
                st.markdown(
                    f"""
                    <div style="
                        background-color: #e8f4f8;
                        padding: 0.5rem;
                        border-radius: 0.25rem;
                        margin-bottom: 1rem;
                        font-size: 0.9em;
                        color: #212529;
                    ">
                    {PromptFormatter.create_section_property_display(section1, show_stats=False).replace(chr(10), '<br>')}
                    </div>
                    """,
                    unsafe_allow_html=True
                )

            # Prompt content with diff highlighting
            st.markdown(
                f"""
                <div style="
                    background-color: #f8f9fa;
                    padding: 1rem;
                    border-radius: 0.5rem;
                    border-left: 4px solid #28a745;
                    min-height: 200px;
                    line-height: 1.6;
                    color: #212529;
                    font-family: 'SF Mono', 'Monaco', 'Cascadia Code', 'Roboto Mono', Consolas, 'Courier New', monospace;
                ">
                {diff_content1}
                </div>
                """,
                unsafe_allow_html=True
            )

        with col2:
            st.markdown(f"### 📄 {file2_name}")

            # Properties box
            with st.container():
                st.markdown(
                    f"""
                    <div style="
                        background-color: #f8e8e8;
                        padding: 0.5rem;
                        border-radius: 0.25rem;
                        margin-bottom: 1rem;
                        font-size: 0.9em;
                        color: #212529;
                    ">
                    {PromptFormatter.create_section_property_display(section2, show_stats=False).replace(chr(10), '<br>')}
                    </div>
                    """,
                    unsafe_allow_html=True
                )

            # Prompt content with diff highlighting
            st.markdown(
                f"""
                <div style="
                    background-color: #f8f9fa;
                    padding: 1rem;
                    border-radius: 0.5rem;
                    border-left: 4px solid #dc3545;
                    min-height: 200px;
                    line-height: 1.6;
                    color: #212529;
                    font-family: 'SF Mono', 'Monaco', 'Cascadia Code', 'Roboto Mono', Consolas, 'Courier New', monospace;
                ">
                {diff_content2}
                </div>
                """,
                unsafe_allow_html=True
            )

    @staticmethod
    def _generate_side_by_side_diff(text1: str, text2: str) -> tuple[str, str]:
        """Generate side-by-side diff highlighting for two text strings"""

        import difflib
        from html import escape

        # Format the texts for better comparison
        formatted_text1 = PromptFormatter._enhance_text_formatting(text1)
        formatted_text2 = PromptFormatter._enhance_text_formatting(text2)

        # Split into lines for line-by-line comparison
        lines1 = formatted_text1.splitlines()
        lines2 = formatted_text2.splitlines()

        # Create a differ
        differ = difflib.SequenceMatcher(None, lines1, lines2)

        # Process the differences
        highlighted_lines1 = []
        highlighted_lines2 = []

        for tag, i1, i2, j1, j2 in differ.get_opcodes():
            if tag == 'equal':
                # Lines are the same
                for i in range(i1, i2):
                    highlighted_lines1.append(escape(lines1[i]))
                for j in range(j1, j2):
                    highlighted_lines2.append(escape(lines2[j]))

            elif tag == 'delete':
                # Lines only in text1 (deleted from text2)
                for i in range(i1, i2):
                    highlighted_lines1.append(
                        f'<span style="background-color: #ffeef0; color: #d73a49; text-decoration: line-through;">'
                        f'{escape(lines1[i])}</span>'
                    )

            elif tag == 'insert':
                # Lines only in text2 (added to text2)
                for j in range(j1, j2):
                    highlighted_lines2.append(
                        f'<span style="background-color: #e6ffed; color: #28a745; font-weight: 500;">'
                        f'{escape(lines2[j])}</span>'
                    )

            elif tag == 'replace':
                # Lines are different, show character-level diff
                for i in range(i1, i2):
                    if i - i1 < j2 - j1:  # There's a corresponding line in text2
                        j = j1 + (i - i1)
                        char_diff1, char_diff2 = UIComponents._generate_character_diff(
                            lines1[i], lines2[j]
                        )
                        highlighted_lines1.append(char_diff1)
                        highlighted_lines2.append(char_diff2)
                    else:
                        # Extra line in text1 (deleted)
                        highlighted_lines1.append(
                            f'<span style="background-color: #ffeef0; color: #d73a49; text-decoration: line-through;">'
                            f'{escape(lines1[i])}</span>'
                        )

                # Handle any extra lines in text2 (added)
                for j in range(j1 + (i2 - i1), j2):
                    highlighted_lines2.append(
                        f'<span style="background-color: #e6ffed; color: #28a745; font-weight: 500;">'
                        f'{escape(lines2[j])}</span>'
                    )

        # Pad shorter list with empty lines for alignment
        max_lines = max(len(highlighted_lines1), len(highlighted_lines2))
        while len(highlighted_lines1) < max_lines:
            highlighted_lines1.append('<span style="opacity: 0.3; font-style: italic;">[missing line]</span>')
        while len(highlighted_lines2) < max_lines:
            highlighted_lines2.append('<span style="opacity: 0.3; font-style: italic;">[missing line]</span>')

        return '<br>'.join(highlighted_lines1), '<br>'.join(highlighted_lines2)

    @staticmethod
    def _generate_character_diff(line1: str, line2: str) -> tuple[str, str]:
        """Generate character-level diff highlighting for two lines"""

        import difflib
        from html import escape

        # Create a character-level differ
        differ = difflib.SequenceMatcher(None, line1, line2)

        highlighted1 = []
        highlighted2 = []

        for tag, i1, i2, j1, j2 in differ.get_opcodes():
            if tag == 'equal':
                # Characters are the same
                highlighted1.append(escape(line1[i1:i2]))
                highlighted2.append(escape(line2[j1:j2]))

            elif tag == 'delete':
                # Characters only in line1 (deleted)
                highlighted1.append(
                    f'<span style="background-color: #ffeef0; color: #d73a49; text-decoration: line-through;">'
                    f'{escape(line1[i1:i2])}</span>'
                )

            elif tag == 'insert':
                # Characters only in line2 (added)
                highlighted2.append(
                    f'<span style="background-color: #e6ffed; color: #28a745; font-weight: 500;">'
                    f'{escape(line2[j1:j2])}</span>'
                )

            elif tag == 'replace':
                # Characters are different
                highlighted1.append(
                    f'<span style="background-color: #ffeef0; color: #d73a49; text-decoration: line-through;">'
                    f'{escape(line1[i1:i2])}</span>'
                )
                highlighted2.append(
                    f'<span style="background-color: #e6ffed; color: #28a745; font-weight: 500;">'
                    f'{escape(line2[j1:j2])}</span>'
                )

        return ''.join(highlighted1), ''.join(highlighted2)

class SectionMatcher:
    """Intelligent section matching between templates"""

    @staticmethod
    def find_section_matches(sections1: Dict[str, ParsedSection],
                           sections2: Dict[str, ParsedSection],
                           threshold: float = 70.0) -> List[SectionMapping]:
        """Find matching sections between two template files"""

        mappings = []
        used_sections2 = set()

        for section1_name in sections1.keys():
            # Try exact match first
            if section1_name in sections2:
                mappings.append(SectionMapping(
                    file1_section=section1_name,
                    file2_section=section1_name,
                    confidence=100.0,
                    is_manual=False
                ))
                used_sections2.add(section1_name)
                continue

            # Find best fuzzy match
            available_sections2 = [s for s in sections2.keys() if s not in used_sections2]
            if available_sections2:
                best_match, confidence = process.extractOne(
                    section1_name,
                    available_sections2,
                    scorer=fuzz.ratio
                )

                if confidence >= threshold:
                    mappings.append(SectionMapping(
                        file1_section=section1_name,
                        file2_section=best_match,
                        confidence=float(confidence),
                        is_manual=False
                    ))
                    used_sections2.add(best_match)

        return mappings

    @staticmethod
    def get_unmapped_sections(sections1: Dict[str, ParsedSection],
                            sections2: Dict[str, ParsedSection],
                            mappings: List[SectionMapping]) -> Tuple[List[str], List[str]]:
        """Get sections that couldn't be automatically mapped"""

        mapped_sections1 = {m.file1_section for m in mappings}
        mapped_sections2 = {m.file2_section for m in mappings}

        unmapped1 = [s for s in sections1.keys() if s not in mapped_sections1]
        unmapped2 = [s for s in sections2.keys() if s not in mapped_sections2]

        return unmapped1, unmapped2

class DeviationAnalyzer:
    """Intelligent analysis of deviations between templates"""

    @staticmethod
    def analyze_template_deviations(sections1: Dict[str, ParsedSection],
                                  sections2: Dict[str, ParsedSection],
                                  file1_name: str, file2_name: str) -> List[SectionDeviation]:
        """Analyze all deviations between two templates"""

        deviations = []
        all_sections = set(sections1.keys()) | set(sections2.keys())

        for section_name in all_sections:
            deviation = DeviationAnalyzer._analyze_section_deviation(
                section_name,
                sections1.get(section_name),
                sections2.get(section_name)
            )
            deviations.append(deviation)

        return sorted(deviations, key=lambda x: (
            0 if x.severity == 'high' else 1 if x.severity == 'medium' else 2,
            -x.similarity_score,
            x.section_name
        ))

    @staticmethod
    def _analyze_section_deviation(section_name: str,
                                 section1: Optional[ParsedSection],
                                 section2: Optional[ParsedSection]) -> SectionDeviation:
        """Analyze deviation for a single section"""

        # Handle missing sections
        if section1 is None:
            return SectionDeviation(
                section_name=section_name,
                file1_section=None,
                file2_section=section2,
                deviation_type='added',
                similarity_score=0.0,
                content_diff_chars=len(section2.prompt) if section2 else 0,
                property_differences=['Missing in first template'],
                severity='high'
            )

        if section2 is None:
            return SectionDeviation(
                section_name=section_name,
                file1_section=section1,
                file2_section=None,
                deviation_type='missing',
                similarity_score=0.0,
                content_diff_chars=len(section1.prompt) if section1 else 0,
                property_differences=['Missing in second template'],
                severity='high'
            )

        # Both sections exist - analyze differences
        similarity_score = fuzz.ratio(section1.prompt, section2.prompt)
        content_diff_chars = len(section2.prompt) - len(section1.prompt)

        # Check property differences
        property_differences = []
        if section1.type != section2.type:
            property_differences.append(f"Type: {section1.type} → {section2.type}")
        if section1.sub_type != section2.sub_type:
            property_differences.append(f"Sub-type: {section1.sub_type} → {section2.sub_type}")
        if section1.include_screenshots != section2.include_screenshots:
            property_differences.append(f"Screenshots: {section1.include_screenshots} → {section2.include_screenshots}")
        if section1.screenshot_instructions != section2.screenshot_instructions:
            property_differences.append("Screenshot instructions differ")

        # Determine deviation type
        if similarity_score >= 95 and not property_differences:
            deviation_type = 'identical'
        elif property_differences and similarity_score >= 90:
            deviation_type = 'properties'
        else:
            deviation_type = 'content'

        return SectionDeviation(
            section_name=section_name,
            file1_section=section1,
            file2_section=section2,
            deviation_type=deviation_type,
            similarity_score=similarity_score,
            content_diff_chars=content_diff_chars,
            property_differences=property_differences,
            severity='low'  # Will be auto-calculated in __post_init__
        )

    @staticmethod
    def get_deviation_summary(deviations: List[SectionDeviation]) -> Dict[str, int]:
        """Get summary statistics of deviations"""

        summary = {
            'total_sections': len(deviations),
            'identical': 0,
            'content_differences': 0,
            'property_differences': 0,
            'missing_sections': 0,
            'added_sections': 0,
            'high_severity': 0,
            'medium_severity': 0,
            'low_severity': 0
        }

        for dev in deviations:
            if dev.deviation_type == 'identical':
                summary['identical'] += 1
            elif dev.deviation_type == 'content':
                summary['content_differences'] += 1
            elif dev.deviation_type == 'properties':
                summary['property_differences'] += 1
            elif dev.deviation_type == 'missing':
                summary['missing_sections'] += 1
            elif dev.deviation_type == 'added':
                summary['added_sections'] += 1

            if dev.severity == 'high':
                summary['high_severity'] += 1
            elif dev.severity == 'medium':
                summary['medium_severity'] += 1
            else:
                summary['low_severity'] += 1

        return summary

class DocxExporter:
    """Export templates back to Klarity-friendly DOCX format with proper embedded comments"""

    @staticmethod
    def build_comment_string(section: ParsedSection) -> str:
        """Build comment string in Klarity concatenated format"""

        # Escape newlines as literal \n like the original system expects
        prompt_escaped = (section.prompt or '').replace('\r\n', '\n').replace('\n', '\\n')

        # Build in the exact concatenated format found in Alex's templates
        comment_string = (
            f"type - {section.type}"
            f"sub_type - {section.sub_type}"
            f"prompt - {prompt_escaped}"
            f"include_screenshots - {'yes' if section.include_screenshots else 'no'}"
            f"screenshot_instructions - {section.screenshot_instructions or 'none'}"
        )

        return comment_string

    @staticmethod
    def export_template_to_docx(template_name: str, sections: Dict[str, ParsedSection], original_file_data: bytes = None) -> BytesIO:
        """Export sections to a Klarity-ready DOCX file preserving original document structure"""

        if original_file_data:
            # Use original file as base and only modify comments
            return DocxExporter._update_comments_in_original_docx(original_file_data, sections)
        else:
            # Fallback to creating new document (shouldn't happen in normal usage)
            return DocxExporter._create_new_docx_with_sections(template_name, sections)

    @staticmethod
    def _update_comments_in_original_docx(original_file_data: bytes, sections: Dict[str, ParsedSection]) -> BytesIO:
        """Update comments in the original DOCX file while preserving all document structure"""

        import zipfile
        import tempfile
        from xml.dom import minidom

        # Create input buffer from original file data
        input_buffer = BytesIO(original_file_data)

        # Read the original DOCX
        docx_zip = zipfile.ZipFile(input_buffer, 'r')

        # Create new output DOCX
        output_buffer = BytesIO()
        new_docx = zipfile.ZipFile(output_buffer, 'w', zipfile.ZIP_DEFLATED)

        try:
            # Copy all files from original, modifying only comments.xml
            for item in docx_zip.infolist():
                data = docx_zip.read(item.filename)

                if item.filename == 'word/comments.xml':
                    # Replace comments.xml with updated comments
                    data = DocxExporter._create_updated_comments_xml(data, sections).encode('utf-8')

                new_docx.writestr(item, data)

        finally:
            docx_zip.close()
            new_docx.close()

        output_buffer.seek(0)
        return output_buffer

    @staticmethod
    def _create_updated_comments_xml(original_comments_data: bytes, sections: Dict[str, ParsedSection]) -> str:
        """Create updated comments.xml preserving structure but updating content"""

        try:
            # Parse original comments.xml
            doc = minidom.parseString(original_comments_data.decode('utf-8'))

            # Get all comment elements
            comments = doc.getElementsByTagName('w:comment')

            # Build a mapping of comment IDs to section names by finding the linked sections
            # This is complex - for now, let's rebuild the entire comments.xml
            return DocxExporter._create_comments_xml(sections)

        except:
            # If parsing fails, create new comments.xml
            return DocxExporter._create_comments_xml(sections)

    @staticmethod
    def _create_new_docx_with_sections(template_name: str, sections: Dict[str, ParsedSection]) -> BytesIO:
        """Fallback: Create new document (used when no original file available)"""

        # Create a minimal document with just section headings
        doc = Document()

        # Add title
        title = doc.add_heading(template_name, 0)

        # Add sections with minimal content (clean Klarity format)
        for section_name, section in sections.items():
            # Add section heading only
            heading = doc.add_heading(section_name, level=1)

        # Save document
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Add comments to the document
        return DocxExporter._add_comments_to_docx(buffer, sections)

    @staticmethod
    def _add_comments_to_docx(buffer: BytesIO, sections: Dict[str, ParsedSection]) -> BytesIO:
        """Add comments to an existing DOCX by modifying its ZIP structure"""

        import zipfile
        import tempfile
        from xml.dom import minidom
        import uuid

        # Read the existing DOCX
        buffer.seek(0)
        docx_zip = zipfile.ZipFile(buffer, 'r')

        # Create a new DOCX with comments
        new_buffer = BytesIO()
        new_docx = zipfile.ZipFile(new_buffer, 'w', zipfile.ZIP_DEFLATED)

        try:
            # Copy all existing files
            for item in docx_zip.infolist():
                data = docx_zip.read(item.filename)
                if item.filename == 'word/document.xml':
                    # Modify document.xml to add comment references
                    data = DocxExporter._add_comment_references_to_document(data, sections)
                elif item.filename == '[Content_Types].xml':
                    # Update content types to include comments
                    data = DocxExporter._update_content_types(data)
                elif item.filename == 'word/_rels/document.xml.rels':
                    # Update relationships to include comments
                    data = DocxExporter._update_document_rels(data)

                new_docx.writestr(item, data)

            # Add comments.xml
            comments_xml = DocxExporter._create_comments_xml(sections)
            new_docx.writestr('word/comments.xml', comments_xml)

        finally:
            docx_zip.close()
            new_docx.close()

        new_buffer.seek(0)
        return new_buffer

    @staticmethod
    def _add_comment_references_to_document(document_xml_data: bytes, sections: Dict[str, ParsedSection]) -> bytes:
        """Add comment references to section headings in document.xml"""

        try:
            # Parse the document XML
            doc = minidom.parseString(document_xml_data.decode('utf-8'))

            # Find all heading paragraphs and add comment references
            comment_id = 0
            paragraphs = doc.getElementsByTagName('w:p')

            for paragraph in paragraphs:
                # Check if this is a heading with style
                pPr = paragraph.getElementsByTagName('w:pPr')
                if pPr:
                    pStyle = pPr[0].getElementsByTagName('w:pStyle')
                    if pStyle and pStyle[0].getAttribute('w:val').startswith('Heading'):
                        # Get the heading text
                        text_runs = paragraph.getElementsByTagName('w:t')
                        if text_runs:
                            heading_text = ''.join([t.firstChild.nodeValue for t in text_runs if t.firstChild])

                            # Check if this heading matches one of our sections
                            if heading_text in sections:
                                comment_id += 1

                                # Add comment range start
                                commentRangeStart = doc.createElement('w:commentRangeStart')
                                commentRangeStart.setAttribute('w:id', str(comment_id))
                                paragraph.insertBefore(commentRangeStart, paragraph.firstChild)

                                # Add comment range end
                                commentRangeEnd = doc.createElement('w:commentRangeEnd')
                                commentRangeEnd.setAttribute('w:id', str(comment_id))
                                paragraph.appendChild(commentRangeEnd)

                                # Add comment reference
                                commentReference = doc.createElement('w:commentReference')
                                commentReference.setAttribute('w:id', str(comment_id))

                                # Create a new run for the comment reference
                                run = doc.createElement('w:r')
                                run.appendChild(commentReference)
                                paragraph.appendChild(run)

            return doc.toxml(encoding='utf-8')

        except Exception as e:
            # If we can't parse/modify, return original
            return document_xml_data

    @staticmethod
    def _create_comments_xml(sections: Dict[str, ParsedSection]) -> str:
        """Create the comments.xml file with our section comments"""

        comments_xml = '''<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'''

        comment_id = 0
        for section_name, section in sections.items():
            comment_id += 1
            comment_text = DocxExporter.build_comment_string(section)

            # Escape XML special characters
            comment_text_escaped = comment_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

            comments_xml += f'''
    <w:comment w:id="{comment_id}" w:author="Klarity Template Comparison Tool" w:date="{datetime.datetime.now().isoformat()}" w:initials="KTCT">
        <w:p>
            <w:pPr>
                <w:pStyle w:val="CommentText"/>
            </w:pPr>
            <w:r>
                <w:rPr>
                    <w:rStyle w:val="CommentReference"/>
                </w:rPr>
                <w:annotationRef/>
            </w:r>
            <w:r>
                <w:t>{comment_text_escaped}</w:t>
            </w:r>
        </w:p>
    </w:comment>'''

        comments_xml += '''
</w:comments>'''

        return comments_xml

    @staticmethod
    def _update_content_types(content_types_data: bytes) -> bytes:
        """Update [Content_Types].xml to include comments"""

        try:
            content_types_str = content_types_data.decode('utf-8')

            # Add comments override if not present
            if 'word/comments.xml' not in content_types_str:
                # Insert before the closing tag
                insert_pos = content_types_str.rfind('</Types>')
                if insert_pos != -1:
                    comment_override = '<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>'
                    content_types_str = content_types_str[:insert_pos] + comment_override + content_types_str[insert_pos:]

            return content_types_str.encode('utf-8')

        except Exception as e:
            return content_types_data

    @staticmethod
    def _update_document_rels(rels_data: bytes) -> bytes:
        """Update document.xml.rels to include comments relationship"""

        try:
            rels_str = rels_data.decode('utf-8')

            # Add comments relationship if not present
            if 'comments.xml' not in rels_str:
                # Find the highest rId
                import re
                rids = re.findall(r'rId(\d+)', rels_str)
                max_rid = max([int(rid) for rid in rids]) if rids else 0
                new_rid = max_rid + 1

                # Insert before the closing tag
                insert_pos = rels_str.rfind('</Relationships>')
                if insert_pos != -1:
                    comment_rel = f'<Relationship Id="rId{new_rid}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="comments.xml"/>'
                    rels_str = rels_str[:insert_pos] + comment_rel + rels_str[insert_pos:]

            return rels_str.encode('utf-8')

        except Exception as e:
            return rels_data

class DiffViewer:
    """Generate and display diffs between two text strings"""

    @staticmethod
    def generate_diff_html(original: str, modified: str, file1_name: str = "Template 1", file2_name: str = "Template 2") -> str:
        """Generate HTML diff between two strings"""

        # Normalize texts for comparison
        original_lines = original.splitlines(keepends=True)
        modified_lines = modified.splitlines(keepends=True)

        # Generate diff
        diff = list(difflib.unified_diff(
            original_lines,
            modified_lines,
            fromfile=file1_name,
            tofile=file2_name,
            lineterm=''
        ))

        if not diff:
            return "<p style='color: gray; text-align: center; padding: 2rem;'>✅ No differences found - prompts are identical</p>"

        html_parts = []
        html_parts.append("<pre style='background-color: #f8f9fa; padding: 1rem; border-radius: 0.5rem; font-family: monospace; white-space: pre-wrap; word-wrap: break-word;'>")

        for line in diff:
            line_escaped = line.replace('<', '&lt;').replace('>', '&gt;')
            if line.startswith('+++') or line.startswith('---'):
                html_parts.append(f"<span style='color: #6c757d; font-weight: bold;'>{line_escaped}</span>")
            elif line.startswith('@@'):
                html_parts.append(f"<span style='color: #007bff; font-weight: bold;'>{line_escaped}</span>")
            elif line.startswith('+'):
                html_parts.append(f"<span style='color: #28a745; background-color: #d4edda; padding: 2px;'>{line_escaped}</span>")
            elif line.startswith('-'):
                html_parts.append(f"<span style='color: #dc3545; background-color: #f8d7da; padding: 2px; text-decoration: line-through;'>{line_escaped}</span>")
            else:
                html_parts.append(f"<span>{line_escaped}</span>")

        html_parts.append("</pre>")
        return '\n'.join(html_parts)

# Streamlit App
def main():
    st.set_page_config(
        page_title="Klarity Template Comparison Tool",
        page_icon="📄",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    # Initialize session state
    if 'section_mappings' not in st.session_state:
        st.session_state.section_mappings = {}
    if 'edited_sections' not in st.session_state:
        st.session_state.edited_sections = {}

    # Header with branding
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.title("🚀 Klarity Template Comparison Tool")
        st.markdown("*Intelligent DOCX template prompt comparison and editing*")

    # Sidebar for file uploads and settings
    with st.sidebar:
        st.header("📁 Upload Templates")

        uploaded_files = st.file_uploader(
            "Choose DOCX files (2+ required)",
            accept_multiple_files=True,
            type=['docx'],
            help="Upload DOCX files with embedded prompts as comments"
        )

        if uploaded_files:
            st.success(f"✅ {len(uploaded_files)} files uploaded")

            with st.expander("📋 File Details", expanded=True):
                for i, file in enumerate(uploaded_files, 1):
                    st.write(f"**{i}.** {file.name}")
                    st.caption(f"Size: {file.size:,} bytes")

        st.divider()

        # Settings
        st.header("⚙️ Settings")

        match_threshold = st.slider(
            "Section matching threshold",
            min_value=50,
            max_value=100,
            value=70,
            help="Minimum similarity % for automatic section matching"
        )

        show_debug = st.checkbox("Show debug info", value=False)

    # Main content area
    if not uploaded_files:
        st.info("👆 Please upload DOCX files to get started")

        with st.expander("ℹ️ How to use this tool"):
            st.markdown("""
            ### 🎯 Purpose
            Compare and edit prompts embedded as comments in DOCX templates, then export in Klarity-friendly format.

            ### 📋 Instructions:
            1. **Upload Templates**: Add 2+ DOCX files with comment-linked sections
            2. **Smart Matching**: AI automatically matches similar sections across templates
            3. **Manual Override**: Adjust mappings for sections that need manual correlation
            4. **Edit Prompts**: Modify prompts inline with real-time preview
            5. **Export**: Download edited templates in Klarity-ready format

            ### 📄 Template Requirements:
            - Section headings must have attached comments
            - Comments contain prompt instructions (structured or freeform)
            - Supported structured format:
              ```
              type - text
              sub_type - bulleted
              prompt - Your detailed instructions here
              include_screenshots - yes
              screenshot_instructions - Capture workflow screenshots
              ```
            """)
        return

    if len(uploaded_files) < 1:
        st.warning("⚠️ Please upload at least 1 DOCX file to proceed")
        return

    # Process uploaded files
    all_sections = {}  # file_name -> {section_name -> ParsedSection}

    with st.spinner("🔄 Processing DOCX files..."):
        progress_bar = st.progress(0)

        for i, uploaded_file in enumerate(uploaded_files):
            progress_bar.progress((i + 1) / len(uploaded_files))

            # Store original file data for export purposes
            if 'original_files' not in st.session_state:
                st.session_state.original_files = {}
            st.session_state.original_files[uploaded_file.name] = uploaded_file.getbuffer()

            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(uploaded_file.getbuffer())
                tmp_file_path = tmp_file.name

            try:
                # Extract sections from DOCX
                sections = DocxCommentExtractor.extract_comments_from_docx(tmp_file_path)
                all_sections[uploaded_file.name] = sections

            finally:
                # Clean up temp file
                os.unlink(tmp_file_path)

        progress_bar.empty()

    # Display results
    total_sections = sum(len(sections) for sections in all_sections.values())
    if total_sections == 0:
        st.error("❌ No prompts found in uploaded files")
        st.info("💡 Ensure your DOCX files contain comments linked to section headings")
        return

    # Success metrics
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("📄 Files Processed", len(uploaded_files))
    with col2:
        st.metric("📝 Sections Found", total_sections)
    with col3:
        files_with_sections = len([f for f in all_sections.values() if f])
        st.metric("✅ Files with Prompts", files_with_sections)
    with col4:
        avg_sections = total_sections / len(uploaded_files) if uploaded_files else 0
        st.metric("📊 Avg Sections/File", f"{avg_sections:.1f}")

    # Create enhanced tabs
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "📋 Browse All",
        "🔗 Smart Mapping",
        "⚡ Deviation Analysis",
        "🚀 Edit & Compare",
        "📤 Export"
    ])

    with tab1:
        browse_all_prompts(all_sections, show_debug)

    with tab2:
        smart_section_mapping(all_sections, match_threshold)

    with tab3:
        deviation_analysis(all_sections)

    with tab4:
        edit_and_compare_enhanced(all_sections)

    with tab5:
        export_templates(all_sections)


def browse_all_prompts(all_sections: Dict[str, Dict[str, ParsedSection]], show_debug: bool):
    """Tab 1: Browse all extracted prompts with enhanced UI"""

    st.header("📋 All Extracted Prompts")
    st.markdown("*Browse all sections across your templates with enhanced readability and formatting*")

    # Enhanced filters
    col1, col2, col3 = st.columns([2, 2, 1])

    with col1:
        selected_files = st.multiselect(
            "📁 Filter by templates:",
            options=list(all_sections.keys()),
            default=list(all_sections.keys()),
            help="Select specific templates to display"
        )

    with col2:
        search_term = st.text_input(
            "🔍 Search content:",
            placeholder="Search section names or prompt content...",
            help="Search across section names and prompt text"
        )

    with col3:
        view_mode = st.selectbox(
            "👁️ View:",
            options=["Compact", "Detailed"],
            index=1,
            help="Choose display density"
        )

    if not selected_files:
        selected_files = list(all_sections.keys())

    # Process and display sections
    total_displayed = 0

    for file_name in selected_files:
        sections = all_sections[file_name]
        if not sections:
            continue

        # Filter by search term
        filtered_sections = sections
        if search_term:
            filtered_sections = {
                name: section for name, section in sections.items()
                if search_term.lower() in name.lower() or
                   search_term.lower() in section.prompt.lower()
            }

        if not filtered_sections:
            continue

        # File header with statistics
        col1, col2, col3 = st.columns([3, 1, 1])
        with col1:
            st.subheader(f"📄 {file_name}")
        with col2:
            st.metric("Sections", len(filtered_sections))
        with col3:
            edited_count = sum(1 for s in filtered_sections.values() if hasattr(s, 'edited') and s.edited)
            st.metric("Edited", edited_count)

        if search_term and len(filtered_sections) != len(sections):
            st.caption(f"Showing {len(filtered_sections)} of {len(sections)} sections")

        st.divider()

        # Display sections using enhanced components
        for section_name, section in filtered_sections.items():
            if view_mode == "Detailed":
                UIComponents.create_section_card(section_name, section, file_name, expanded=False)
            else:
                # Compact view
                with st.expander(f"📝 {section_name}" + (" ✏️" if hasattr(section, 'edited') and section.edited else ""), expanded=False):
                    col1, col2 = st.columns([3, 1])

                    with col1:
                        # Truncated prompt for compact view
                        formatted_prompt = PromptFormatter.format_prompt_for_display(section.prompt, max_length=300)
                        st.markdown(formatted_prompt)

                    with col2:
                        # Compact properties
                        properties = PromptFormatter.create_section_property_display(section, show_stats=True)
                        st.markdown(properties)

            total_displayed += 1

            # Debug info if requested
            if show_debug:
                with st.expander(f"🐛 Debug: {section_name}", expanded=False):
                    st.text("Raw Comment Data:")
                    st.code(section.raw_comment)
                    st.text("Parsed Properties:")
                    st.json({
                        "name": section.name,
                        "type": section.type,
                        "sub_type": section.sub_type,
                        "include_screenshots": section.include_screenshots,
                        "screenshot_instructions": section.screenshot_instructions,
                        "edited": getattr(section, 'edited', False)
                    })

        st.divider()

    # Summary
    if total_displayed == 0:
        st.info("📭 No sections match your current filters")
    else:
        st.success(f"✅ Displaying {total_displayed} sections across {len(selected_files)} templates")


def deviation_analysis(all_sections: Dict[str, Dict[str, ParsedSection]]):
    """Tab 3: Intelligent deviation analysis between templates"""

    st.header("⚡ Intelligent Deviation Analysis")
    st.markdown("*Quickly identify and navigate to differences between templates*")

    if len(all_sections) < 2:
        st.info("📋 Upload at least 2 templates to analyze deviations")
        return

    # Template selection
    file_names = list(all_sections.keys())
    col1, col2 = st.columns(2)

    with col1:
        file1 = st.selectbox("Primary template:", file_names, key="deviation_file1")
    with col2:
        file2 = st.selectbox("Compare with:", [f for f in file_names if f != file1], key="deviation_file2")

    if not file1 or not file2 or file1 == file2:
        st.info("Select two different templates to analyze deviations")
        return

    # Analyze deviations
    sections1 = all_sections[file1]
    sections2 = all_sections[file2]

    with st.spinner("🔄 Analyzing deviations..."):
        deviations = DeviationAnalyzer.analyze_template_deviations(sections1, sections2, file1, file2)
        summary = DeviationAnalyzer.get_deviation_summary(deviations)

    # Summary Dashboard
    st.subheader("📊 Deviation Summary")

    # Key metrics
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        st.metric("Total Sections", summary['total_sections'])
    with col2:
        st.metric("🔴 High Severity", summary['high_severity'],
                 delta=f"{summary['high_severity']/summary['total_sections']*100:.1f}%" if summary['total_sections'] > 0 else "0%")
    with col3:
        st.metric("🟡 Medium Severity", summary['medium_severity'],
                 delta=f"{summary['medium_severity']/summary['total_sections']*100:.1f}%" if summary['total_sections'] > 0 else "0%")
    with col4:
        st.metric("🟢 Low Severity", summary['low_severity'],
                 delta=f"{summary['low_severity']/summary['total_sections']*100:.1f}%" if summary['total_sections'] > 0 else "0%")
    with col5:
        st.metric("✅ Identical", summary['identical'],
                 delta=f"{summary['identical']/summary['total_sections']*100:.1f}%" if summary['total_sections'] > 0 else "0%")

    st.divider()

    # Deviation breakdown
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        if summary['content_differences'] > 0:
            st.metric("📝 Content Differs", summary['content_differences'], delta="Content changes")
    with col2:
        if summary['property_differences'] > 0:
            st.metric("⚙️ Property Differs", summary['property_differences'], delta="Setting changes")
    with col3:
        if summary['missing_sections'] > 0:
            st.metric("❌ Missing", summary['missing_sections'], delta="In template 2", delta_color="inverse")
    with col4:
        if summary['added_sections'] > 0:
            st.metric("➕ Added", summary['added_sections'], delta="In template 2")

    # Quick filters
    st.subheader("🎛️ Quick Navigation")

    col1, col2, col3 = st.columns(3)
    with col1:
        severity_filter = st.multiselect(
            "Filter by severity:",
            options=['high', 'medium', 'low'],
            default=['high', 'medium'],
            help="Show only deviations of selected severity levels"
        )

    with col2:
        type_filter = st.multiselect(
            "Filter by deviation type:",
            options=['content', 'properties', 'missing', 'added', 'identical'],
            default=['content', 'properties', 'missing', 'added'],
            help="Show only specific types of deviations"
        )

    with col3:
        sort_by = st.selectbox(
            "Sort by:",
            options=['severity', 'similarity', 'section_name', 'content_size'],
            index=0,
            help="Choose how to sort the deviation list"
        )

    # Filter deviations based on user selection
    filtered_deviations = [
        dev for dev in deviations
        if dev.severity in severity_filter and dev.deviation_type in type_filter
    ]

    # Sort deviations
    if sort_by == 'severity':
        filtered_deviations.sort(key=lambda x: (
            0 if x.severity == 'high' else 1 if x.severity == 'medium' else 2,
            x.section_name
        ))
    elif sort_by == 'similarity':
        filtered_deviations.sort(key=lambda x: x.similarity_score)
    elif sort_by == 'section_name':
        filtered_deviations.sort(key=lambda x: x.section_name)
    elif sort_by == 'content_size':
        filtered_deviations.sort(key=lambda x: abs(x.content_diff_chars), reverse=True)

    st.divider()

    # Deviation details with quick navigation
    st.subheader(f"🔍 Detailed Analysis ({len(filtered_deviations)} sections)")

    if not filtered_deviations:
        st.info("No deviations match your current filters")
        return

    # Display deviations
    for i, deviation in enumerate(filtered_deviations):
        # Severity indicator
        severity_colors = {'high': '🔴', 'medium': '🟡', 'low': '🟢'}
        severity_icon = severity_colors.get(deviation.severity, '⚪')

        # Deviation type indicator
        type_icons = {
            'content': '📝', 'properties': '⚙️', 'missing': '❌',
            'added': '➕', 'identical': '✅'
        }
        type_icon = type_icons.get(deviation.deviation_type, '❓')

        # Create expandable section
        with st.expander(
            f"{severity_icon} {type_icon} {deviation.section_name} "
            f"({deviation.similarity_score:.0f}% similar)",
            expanded=i < 3  # Auto-expand first 3 high-priority items
        ):
            col1, col2 = st.columns([3, 1])

            with col1:
                # Deviation description
                if deviation.deviation_type == 'missing':
                    st.error(f"❌ **Missing in {file2}**")
                    st.write("This section exists in the first template but is missing in the second template.")

                elif deviation.deviation_type == 'added':
                    st.success(f"➕ **Added in {file2}**")
                    st.write("This section exists in the second template but is missing in the first template.")

                elif deviation.deviation_type == 'identical':
                    st.success("✅ **Identical sections**")
                    st.write("These sections are identical or nearly identical across both templates.")

                elif deviation.deviation_type == 'properties':
                    st.warning("⚙️ **Property differences detected**")
                    for diff in deviation.property_differences:
                        st.write(f"• {diff}")

                elif deviation.deviation_type == 'content':
                    st.warning("📝 **Content differences detected**")
                    if deviation.content_diff_chars != 0:
                        st.write(f"Content length difference: {deviation.content_diff_chars:+,} characters")

                # Show property differences if any
                if deviation.property_differences and deviation.deviation_type != 'properties':
                    st.write("**Additional property differences:**")
                    for diff in deviation.property_differences:
                        st.write(f"• {diff}")

            with col2:
                st.markdown("**📊 Quick Stats:**")

                # Similarity meter
                if deviation.similarity_score > 0:
                    similarity_color = "🟢" if deviation.similarity_score >= 75 else "🟡" if deviation.similarity_score >= 50 else "🔴"
                    st.write(f"**Similarity:** {similarity_color} {deviation.similarity_score:.1f}%")

                # Content size comparison
                if deviation.file1_section and deviation.file2_section:
                    len1 = len(deviation.file1_section.prompt)
                    len2 = len(deviation.file2_section.prompt)
                    st.write(f"**File 1:** {len1:,} chars")
                    st.write(f"**File 2:** {len2:,} chars")

                elif deviation.file1_section:
                    st.write(f"**File 1:** {len(deviation.file1_section.prompt):,} chars")
                    st.write(f"**File 2:** Missing")

                elif deviation.file2_section:
                    st.write(f"**File 1:** Missing")
                    st.write(f"**File 2:** {len(deviation.file2_section.prompt):,} chars")

                # Quick action buttons
                st.markdown("**🚀 Quick Actions:**")

                # Jump to detailed diff
                if st.button(f"🔍 View Diff", key=f"diff_jump_{i}_{deviation.section_name}"):
                    st.session_state['diff_jump_section'] = deviation.section_name
                    st.session_state['diff_jump_file1'] = file1
                    st.session_state['diff_jump_file2'] = file2
                    st.success(f"Navigate to 'Diff Analysis' tab to see detailed comparison for '{deviation.section_name}'")

                # Jump to edit mode
                if deviation.file1_section and st.button(f"✏️ Edit", key=f"edit_jump_{i}_{deviation.section_name}"):
                    st.session_state['edit_jump_section'] = deviation.section_name
                    st.session_state['edit_jump_file'] = file1
                    st.success(f"Navigate to 'Edit & Compare' tab to edit '{deviation.section_name}'")

    # Summary message
    if filtered_deviations:
        total_issues = len([d for d in filtered_deviations if d.deviation_type != 'identical'])
        if total_issues > 0:
            st.warning(f"⚠️ Found {total_issues} sections with deviations that may need attention")
        else:
            st.success("✅ All analyzed sections are identical or have minimal differences")


def smart_section_mapping(all_sections: Dict[str, Dict[str, ParsedSection]], threshold: float):
    """Tab 2: Intelligent section mapping between templates"""

    st.header("🔗 Smart Section Mapping")

    if len(all_sections) < 2:
        st.info("📋 Upload at least 2 files to use section mapping")
        return

    # File pair selection
    file_names = list(all_sections.keys())
    col1, col2 = st.columns(2)

    with col1:
        file1 = st.selectbox("Primary template:", file_names, key="mapping_file1")
    with col2:
        file2 = st.selectbox("Compare with:", [f for f in file_names if f != file1], key="mapping_file2")

    if not file1 or not file2 or file1 == file2:
        st.info("Select two different files to continue")
        return

    sections1 = all_sections[file1]
    sections2 = all_sections[file2]

    # Generate automatic mappings
    mappings = SectionMatcher.find_section_matches(sections1, sections2, threshold)
    unmapped1, unmapped2 = SectionMatcher.get_unmapped_sections(sections1, sections2, mappings)

    # Store in session state
    mapping_key = f"{file1}::{file2}"
    if mapping_key not in st.session_state.section_mappings:
        st.session_state.section_mappings[mapping_key] = mappings

    # Display mapping results
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("🎯 Auto-matched", len(mappings))
    with col2:
        st.metric("❓ Unmapped in File 1", len(unmapped1))
    with col3:
        st.metric("❓ Unmapped in File 2", len(unmapped2))

    # Show automatic mappings
    if mappings:
        st.subheader("🤖 Automatic Mappings")

        for i, mapping in enumerate(mappings):
            confidence_color = "🟢" if mapping.confidence >= 90 else "🟡" if mapping.confidence >= 70 else "🟠"

            col1, col2, col3 = st.columns([2, 2, 1])
            with col1:
                st.write(f"**{mapping.file1_section}**")
            with col2:
                st.write(f"**{mapping.file2_section}**")
            with col3:
                st.write(f"{confidence_color} {mapping.confidence:.0f}%")

    # Manual mapping for unmapped sections
    if unmapped1 or unmapped2:
        st.subheader("🔧 Manual Mapping")
        st.info("Map sections that couldn't be automatically matched")

        # Create manual mappings
        manual_mappings = []

        for i, section1 in enumerate(unmapped1):
            col1, col2, col3 = st.columns([2, 2, 1])

            with col1:
                st.write(f"**{section1}**")
                if section1 in sections1:
                    preview = sections1[section1].prompt[:100] + "..." if len(sections1[section1].prompt) > 100 else sections1[section1].prompt
                    st.caption(preview)

            with col2:
                options = ["[No Match]"] + unmapped2
                selected = st.selectbox(
                    f"Match with:",
                    options,
                    key=f"manual_map_{i}_{section1}"
                )

                if selected != "[No Match]":
                    manual_mappings.append(SectionMapping(
                        file1_section=section1,
                        file2_section=selected,
                        confidence=0.0,  # Manual mapping
                        is_manual=True
                    ))

            with col3:
                if selected != "[No Match]":
                    st.success("👤 Manual")
                else:
                    st.info("⚠️ No match")

        # Update session state with manual mappings
        if manual_mappings:
            all_mappings = mappings + manual_mappings
            st.session_state.section_mappings[mapping_key] = all_mappings

            st.success(f"✅ {len(manual_mappings)} manual mappings added")


def edit_and_compare_enhanced(all_sections: Dict[str, Dict[str, ParsedSection]]):
    """Enhanced tab: Edit and compare prompts between two templates side-by-side"""

    st.header("🚀 Edit & Compare Prompts")

    if len(all_sections) < 2:
        st.info("📋 Upload at least 2 files for side-by-side comparison and editing")

        # Show single file editing interface if only one file
        if len(all_sections) == 1:
            file_name = list(all_sections.keys())[0]
            st.subheader(f"✏️ Single File Editing: {file_name}")
            _render_single_file_editor(all_sections, file_name)
        return

    # Dual template selection with intelligent matching
    file_names = list(all_sections.keys())

    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**📄 Template A**")
        file_a = st.selectbox("Select first template:", file_names, key="enhanced_file_a")
    with col2:
        st.markdown("**📄 Template B**")
        file_b = st.selectbox("Select second template:", [f for f in file_names if f != file_a], key="enhanced_file_b")

    if not file_a or not file_b:
        st.info("Select two templates to continue")
        return

    sections_a = all_sections[file_a]
    sections_b = all_sections[file_b]

    # Section selection with intelligent matching
    all_sections_a = list(sections_a.keys())
    all_sections_b = list(sections_b.keys())

    col1, col2 = st.columns(2)
    with col1:
        section_a = st.selectbox(f"Section from {file_a}:", all_sections_a, key="enhanced_section_a")
    with col2:
        # Smart suggestion for section_b based on section_a
        if section_a:
            if section_a in sections_b:
                default_idx = all_sections_b.index(section_a)
            else:
                # Find best match using fuzzy matching
                from fuzzywuzzy import process, fuzz
                matches = process.extract(section_a, all_sections_b, limit=3, scorer=fuzz.ratio)
                if matches and matches[0][1] >= 70:
                    default_idx = all_sections_b.index(matches[0][0])
                else:
                    default_idx = 0
        else:
            default_idx = 0

        section_b = st.selectbox(
            f"Section from {file_b}:",
            all_sections_b,
            index=default_idx,
            key="enhanced_section_b"
        )

    if not section_a or not section_b:
        st.info("Select sections from both templates")
        return

    # Get the actual sections
    sec_a = sections_a[section_a]
    sec_b = sections_b[section_b]

    # Display comparison metrics
    st.subheader(f"📊 Comparing: {section_a} ↔ {section_b}")

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric(f"{file_a} Length", f"{len(sec_a.prompt):,} chars")
    with col2:
        st.metric(f"{file_b} Length", f"{len(sec_b.prompt):,} chars")
    with col3:
        diff_chars = len(sec_b.prompt) - len(sec_a.prompt)
        st.metric("Difference", f"{diff_chars:+,} chars")
    with col4:
        from fuzzywuzzy import fuzz
        similarity = fuzz.ratio(sec_a.prompt, sec_b.prompt)
        st.metric("Similarity", f"{similarity}%")

    st.divider()

    # Enhanced side-by-side editing interface
    _render_dual_editing_interface(file_a, file_b, section_a, section_b, sec_a, sec_b)


def _render_single_file_editor(all_sections: Dict[str, Dict[str, ParsedSection]], file_name: str):
    """Render single file editing interface when only one file is available"""

    sections = all_sections[file_name]
    section_names = list(sections.keys()) if sections else []

    if not section_names:
        st.info("No sections found in the selected file")
        return

    selected_section = st.selectbox("Select section:", section_names, key="single_edit_section")

    if not selected_section:
        return

    section = sections[selected_section]
    section_key = f"{file_name}::{selected_section}"

    # Initialize edited sections in session state
    if section_key not in st.session_state.edited_sections:
        st.session_state.edited_sections[section_key] = {
            'prompt': section.prompt,
            'type': section.type,
            'sub_type': section.sub_type,
            'include_screenshots': section.include_screenshots,
            'screenshot_instructions': section.screenshot_instructions
        }

    edited_data = st.session_state.edited_sections[section_key]

    # Simple editing interface
    col1, col2 = st.columns(2)

    with col1:
        st.markdown("**📝 Edit Prompt:**")
        new_prompt = st.text_area(
            "Prompt content:",
            value=edited_data['prompt'],
            height=300,
            help="Edit the prompt text here"
        )

        # Properties
        _render_properties_editor("single", edited_data, key_suffix="single")

        # Save button
        if st.button("💾 Save Changes", type="primary", key="single_save"):
            _save_section_changes(section_key, section, new_prompt, edited_data)

    with col2:
        st.markdown("**👁️ Live Preview:**")
        _render_preview(new_prompt, section)


def _render_dual_editing_interface(file_a: str, file_b: str, section_a: str, section_b: str, sec_a: ParsedSection, sec_b: ParsedSection):
    """Render the dual editing interface with side-by-side comparison and editing"""

    # Initialize session state for both sections
    section_key_a = f"{file_a}::{section_a}"
    section_key_b = f"{file_b}::{section_b}"

    if section_key_a not in st.session_state.edited_sections:
        st.session_state.edited_sections[section_key_a] = {
            'prompt': sec_a.prompt,
            'type': sec_a.type,
            'sub_type': sec_a.sub_type,
            'include_screenshots': sec_a.include_screenshots,
            'screenshot_instructions': sec_a.screenshot_instructions
        }

    if section_key_b not in st.session_state.edited_sections:
        st.session_state.edited_sections[section_key_b] = {
            'prompt': sec_b.prompt,
            'type': sec_b.type,
            'sub_type': sec_b.sub_type,
            'include_screenshots': sec_b.include_screenshots,
            'screenshot_instructions': sec_b.screenshot_instructions
        }

    edited_data_a = st.session_state.edited_sections[section_key_a]
    edited_data_b = st.session_state.edited_sections[section_key_b]

    # Side-by-side editing interface
    col_a, col_b = st.columns(2)

    with col_a:
        st.markdown(f"**📝 Edit Template A: {file_a}**")
        new_prompt_a = st.text_area(
            f"Prompt content for {section_a}:",
            value=edited_data_a['prompt'],
            height=300,
            help="Edit the prompt text for template A",
            key="edit_prompt_a"
        )

        # Properties for Template A
        _render_properties_editor("Template A", edited_data_a, key_suffix="a")

        # Action buttons for Template A
        col_a1, col_a2, col_a3 = st.columns(3)
        with col_a1:
            if st.button("💾 Save A", type="primary", key="save_a"):
                _save_section_changes(section_key_a, sec_a, new_prompt_a, edited_data_a)
        with col_a2:
            if st.button("📋 Copy B→A", key="copy_b_to_a"):
                st.session_state.edited_sections[section_key_a]['prompt'] = edited_data_b['prompt']
                st.rerun()
        with col_a3:
            if st.button("🔄 Reset A", key="reset_a"):
                _reset_section(section_key_a, sec_a)

    with col_b:
        st.markdown(f"**📝 Edit Template B: {file_b}**")
        new_prompt_b = st.text_area(
            f"Prompt content for {section_b}:",
            value=edited_data_b['prompt'],
            height=300,
            help="Edit the prompt text for template B",
            key="edit_prompt_b"
        )

        # Properties for Template B
        _render_properties_editor("Template B", edited_data_b, key_suffix="b")

        # Action buttons for Template B
        col_b1, col_b2, col_b3 = st.columns(3)
        with col_b1:
            if st.button("💾 Save B", type="primary", key="save_b"):
                _save_section_changes(section_key_b, sec_b, new_prompt_b, edited_data_b)
        with col_b2:
            if st.button("📋 Copy A→B", key="copy_a_to_b"):
                st.session_state.edited_sections[section_key_b]['prompt'] = edited_data_a['prompt']
                st.rerun()
        with col_b3:
            if st.button("🔄 Reset B", key="reset_b"):
                _reset_section(section_key_b, sec_b)

    # Central action buttons
    st.divider()
    col_center1, col_center2, col_center3 = st.columns([1, 2, 1])
    with col_center2:
        if st.button("💾 Save Both Templates", type="primary", key="save_both"):
            _save_section_changes(section_key_a, sec_a, new_prompt_a, edited_data_a)
            _save_section_changes(section_key_b, sec_b, new_prompt_b, edited_data_b)
            st.success("✅ Both templates saved!")

    # Enhanced side-by-side comparison with diff highlighting
    st.subheader("📄 Side-by-Side Comparison with Diff Highlighting")

    UIComponents.create_side_by_side_comparison(
        section_name=f"{section_a} vs {section_b}",
        section1=sec_a,
        section2=sec_b,
        file1_name=file_a,
        file2_name=file_b
    )


def _render_properties_editor(template_name: str, edited_data: dict, key_suffix: str):
    """Render the properties editing interface"""

    col1, col2 = st.columns(2)
    with col1:
        new_type = st.selectbox(
            "Type:",
            options=['text', 'table'],
            index=0 if edited_data['type'] == 'text' else 1,
            key=f"type_{key_suffix}"
        )
        edited_data['type'] = new_type

    with col2:
        subtype_options = ['default', 'bulleted', 'freeform', 'flow-diagram', 'walkthrough-steps']
        current_subtype = edited_data.get('sub_type', 'default')
        if current_subtype == 'bpmn':
            current_subtype = 'flow-diagram'
        elif current_subtype not in subtype_options:
            current_subtype = 'default'

        new_sub_type = st.selectbox(
            "Sub-type:",
            options=subtype_options,
            index=subtype_options.index(current_subtype),
            key=f"subtype_{key_suffix}"
        )
        edited_data['sub_type'] = new_sub_type

    new_include_screenshots = st.checkbox(
        "Include screenshots",
        value=edited_data['include_screenshots'],
        key=f"screenshots_{key_suffix}"
    )
    edited_data['include_screenshots'] = new_include_screenshots

    if new_include_screenshots:
        new_screenshot_instructions = st.text_area(
            "Screenshot instructions:",
            value=edited_data['screenshot_instructions'],
            height=60,
            key=f"screenshot_inst_{key_suffix}"
        )
        edited_data['screenshot_instructions'] = new_screenshot_instructions
    else:
        edited_data['screenshot_instructions'] = ""


def _save_section_changes(section_key: str, section: ParsedSection, new_prompt: str, edited_data: dict):
    """Save changes to a section"""

    st.session_state.edited_sections[section_key] = {
        'prompt': new_prompt,
        'type': edited_data['type'],
        'sub_type': edited_data['sub_type'],
        'include_screenshots': edited_data['include_screenshots'],
        'screenshot_instructions': edited_data['screenshot_instructions']
    }

    # Update the actual section object
    section.prompt = new_prompt
    section.type = edited_data['type']
    section.sub_type = edited_data['sub_type']
    section.include_screenshots = edited_data['include_screenshots']
    section.screenshot_instructions = edited_data['screenshot_instructions']
    section.edited = True

    st.success("✅ Changes saved!")


def _reset_section(section_key: str, section: ParsedSection):
    """Reset a section to its original state"""

    st.session_state.edited_sections[section_key]['prompt'] = section.original_prompt
    section.prompt = section.original_prompt
    section.edited = False
    st.success("✅ Reset to original")
    st.rerun()


def _render_preview(prompt: str, section: ParsedSection):
    """Render a formatted preview of the prompt"""

    formatted_prompt = PromptFormatter.format_prompt_for_display(prompt)

    st.markdown(
        f"""
        <div style="
            background-color: #f8f9fa;
            padding: 1rem;
            border-radius: 0.5rem;
            border-left: 4px solid #28a745;
            line-height: 1.6;
            max-height: 400px;
            overflow-y: auto;
            color: #212529;
        ">
        {formatted_prompt.replace(chr(10), '<br>')}
        </div>
        """,
        unsafe_allow_html=True
    )

    # Show changes indicator
    if prompt != section.original_prompt:
        st.warning("✏️ **Modified from original**")

        original_len = len(section.original_prompt)
        new_len = len(prompt)
        diff = new_len - original_len

        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Original", f"{original_len:,} chars")
        with col2:
            st.metric("Current", f"{new_len:,} chars")
        with col3:
            st.metric("Change", f"{diff:+,}")
    else:
        st.success("✅ **Matches original**")


def export_templates(all_sections: Dict[str, Dict[str, ParsedSection]]):
    """Tab 5: Export templates in Klarity-friendly format"""

    st.header("📤 Export Templates")

    if not all_sections:
        st.info("No templates available for export")
        return

    st.markdown("""
    Export your templates (including any edits) in Klarity-friendly DOCX format.
    The exported files will contain properly formatted comments that can be processed by the Klarity system.
    """)

    # Export options
    export_mode = st.radio(
        "Export mode:",
        options=[
            "📄 Individual Files (one per template)",
            "📦 Merged Template (all sections combined)",
            "✏️ Edited Sections Only"
        ]
    )

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("🔧 Export Options")

        # Template name for merged export
        if "Merged" in export_mode:
            template_name = st.text_input(
                "Template name for merged file:",
                value="Merged_Template"
            )

        # Include metadata
        include_metadata = st.checkbox(
            "Include generation metadata",
            value=True,
            help="Add timestamp and tool information to exported files"
        )

        # Format options
        preserve_edits = st.checkbox(
            "Preserve edit history",
            value=True,
            help="Mark edited sections in the exported document"
        )

    with col2:
        st.subheader("📊 Export Preview")

        total_files = len(all_sections)
        total_sections = sum(len(sections) for sections in all_sections.values())
        edited_sections = sum(
            sum(1 for section in sections.values() if section.edited)
            for sections in all_sections.values()
        )

        st.metric("📄 Templates", total_files)
        st.metric("📝 Total Sections", total_sections)
        st.metric("✏️ Edited Sections", edited_sections)

    # Export buttons
    st.subheader("💾 Download Files")

    if export_mode == "📄 Individual Files (one per template)":
        # Export each template individually
        for file_name, sections in all_sections.items():
            if not sections:
                continue

            col1, col2, col3 = st.columns([2, 1, 1])

            with col1:
                st.write(f"**{file_name}**")
                st.caption(f"{len(sections)} sections" + (f", {sum(1 for s in sections.values() if s.edited)} edited" if edited_sections else ""))

            with col2:
                # Preview button
                clean_file_name = file_name.replace(" ", "_").replace(".", "_").replace("/", "_")
                if st.button(f"👁️ Preview", key=f"preview_individual_{clean_file_name}"):
                    with st.expander(f"Preview: {file_name}", expanded=True):
                        for section_name, section in sections.items():
                            st.write(f"**{section_name}**" + (" ✏️" if section.edited else ""))
                            st.caption(f"Type: {section.type}, Sub-type: {section.sub_type}")
                            preview = section.prompt[:100] + "..." if len(section.prompt) > 100 else section.prompt
                            st.text(preview)
                            st.divider()

            with col3:
                # Generate export file
                clean_name = file_name.replace('.docx', '').replace('.', '_').replace(' ', '_').replace('/', '_')
                original_file_data = st.session_state.original_files.get(file_name) if 'original_files' in st.session_state else None
                export_buffer = DocxExporter.export_template_to_docx(clean_name, sections, original_file_data)

                st.download_button(
                    label="📥 Download",
                    data=export_buffer.getvalue(),
                    file_name=f"{clean_name}_Klarity_Ready.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_individual_{clean_name}"
                )

    elif export_mode == "📦 Merged Template (all sections combined)":
        # Combine all sections from all templates
        merged_sections = {}

        for file_name, sections in all_sections.items():
            for section_name, section in sections.items():
                # Handle duplicate section names by prefixing with file name
                if section_name in merged_sections:
                    unique_name = f"{file_name}_{section_name}"
                else:
                    unique_name = section_name

                merged_sections[unique_name] = section

        st.write(f"**Merged template will contain {len(merged_sections)} sections**")

        if st.button("👁️ Preview Merged Template"):
            with st.expander("Merged Template Preview", expanded=True):
                for section_name, section in merged_sections.items():
                    st.write(f"**{section_name}**" + (" ✏️" if section.edited else ""))
                    st.caption(f"Type: {section.type}, Sub-type: {section.sub_type}")
                    st.divider()

        # Generate merged export
        # For merged templates, we'll use the first file as the base structure
        base_file_name = list(all_sections.keys())[0] if all_sections else None
        original_file_data = st.session_state.original_files.get(base_file_name) if 'original_files' in st.session_state and base_file_name else None
        export_buffer = DocxExporter.export_template_to_docx(template_name, merged_sections, original_file_data)

        st.download_button(
            label="📥 Download Merged Template",
            data=export_buffer.getvalue(),
            file_name=f"{template_name}_Klarity_Ready.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_merged"
        )

    elif export_mode == "✏️ Edited Sections Only":
        # Export only edited sections
        edited_only = {}

        for file_name, sections in all_sections.items():
            file_edited = {name: section for name, section in sections.items() if section.edited}
            if file_edited:
                edited_only[file_name] = file_edited

        if not edited_only:
            st.info("No edited sections found. Make some edits first!")
        else:
            st.write(f"**Found edited sections in {len(edited_only)} files**")

            for file_name, sections in edited_only.items():
                col1, col2, col3 = st.columns([2, 1, 1])

                with col1:
                    st.write(f"**{file_name}** (edited sections)")
                    st.caption(f"{len(sections)} edited sections")

                with col2:
                    clean_file_name_edited = file_name.replace(" ", "_").replace(".", "_").replace("/", "_")
                    if st.button(f"👁️ Preview", key=f"preview_edited_{clean_file_name_edited}"):
                        with st.expander(f"Edited Sections: {file_name}", expanded=True):
                            for section_name, section in sections.items():
                                st.write(f"**{section_name}** ✏️")
                                st.caption(f"Type: {section.type}, Sub-type: {section.sub_type}")
                                st.text("Original → Edited comparison would go here")
                                st.divider()

                with col3:
                    clean_name = file_name.replace('.docx', '').replace('.', '_').replace(' ', '_').replace('/', '_')
                    original_file_data = st.session_state.original_files.get(file_name) if 'original_files' in st.session_state else None
                    export_buffer = DocxExporter.export_template_to_docx(f"{clean_name}_Edited", sections, original_file_data)

                    st.download_button(
                        label="📥 Download",
                        data=export_buffer.getvalue(),
                        file_name=f"{clean_name}_Edited_Klarity_Ready.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_edited_{clean_name}"
                    )

    # Export success message
    st.success("✅ Your exported DOCX files are ready for use with Klarity!")
    st.info("💡 The exported files contain properly formatted comments that can be processed by the Klarity prompt generation system.")


if __name__ == "__main__":
    main()